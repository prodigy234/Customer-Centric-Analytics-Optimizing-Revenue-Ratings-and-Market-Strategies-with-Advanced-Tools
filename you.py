import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import tempfile
import os
import pdfkit

# ---------------- Helper Functions ---------------- #

def add_heading(document, text, level=1):
    heading = document.add_heading(level=level)
    run = heading.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(14 if level == 1 else 12)
    return heading

def add_paragraph(document, text, bold=False):
    para = document.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    return para

def generate_resume(doc, user_data):
    if user_data["photo"]:
        image_path = os.path.join(tempfile.gettempdir(), user_data["photo"].name)
        with open(image_path, "wb") as f:
            f.write(user_data["photo"].read())
        doc.add_picture(image_path, width=Inches(1.5))

    heading = doc.add_heading(user_data["name"], 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, f"📧 {user_data['email']} | 📞 {user_data['phone']}")
    doc.add_paragraph("")

    if user_data["summary"]:
        add_heading(doc, "🎯 Career Summary", level=1)
        add_paragraph(doc, user_data["summary"])

    if user_data["institution"]:
        add_heading(doc, "📚 Education", level=1)
        add_paragraph(doc, f"{user_data['edu_level']} in {user_data['field']} - {user_data['institution']} ({user_data['edu_duration']})")

    if user_data["job_title"]:
        add_heading(doc, "💼 Work Experience", level=1)
        add_paragraph(doc, f"{user_data['job_title']} at {user_data['company']} ({user_data['work_years']}) - {user_data['exp_years']} years")

    if user_data["tech_skills"]:
        add_heading(doc, "🛠 Technical Skills", level=1)
        for skill in user_data["tech_skills"].split(","):
            add_paragraph(doc, f"• {skill.strip()}")

    if user_data["soft_skills"]:
        add_heading(doc, "🤝 Soft Skills", level=1)
        for skill in user_data["soft_skills"].split(","):
            add_paragraph(doc, f"• {skill.strip()}")

    if user_data["certifications"]:
        add_heading(doc, "📜 Certifications", level=1)
        for cert in user_data["certifications"].split(","):
            add_paragraph(doc, f"• {cert.strip()}")

    if user_data["achievements"]:
        add_heading(doc, "🏆 Achievements", level=1)
        for ach in user_data["achievements"].split(","):
            add_paragraph(doc, f"• {ach.strip()}")

    if user_data["projects"]:
        add_heading(doc, "🌐 Projects", level=1)
        for proj in user_data["projects"].split(","):
            add_paragraph(doc, f"• {proj.strip()}")

    if user_data["languages"]:
        add_heading(doc, "🗣️ Languages", level=1)
        for lang in user_data["languages"].split(","):
            add_paragraph(doc, f"• {lang.strip()}")

# ---------------- Streamlit UI ---------------- #

st.set_page_config("AI Resume Builder", layout="centered")
st.markdown("<h1 style='text-align:center; color:#4CAF50;'>📄 Smart AI Resume Builder</h1>", unsafe_allow_html=True)
st.write("Customize your resume with professional templates and download in both DOCX and PDF formats.")

template_style = st.selectbox("Choose a Template Style", ["Classic", "Modern", "Creative"])

st.markdown("---")
st.subheader("👤 Personal Information")

name = st.text_input("Full Name")
email = st.text_input("Email")
phone = st.text_input("Phone Number")
photo = st.file_uploader("Upload a Profile Photo", type=["jpg", "jpeg", "png"])

st.subheader("🎯 Career Summary")
summary = st.text_area("Summary")

st.subheader("📚 Education")
edu_level = st.selectbox("Education Level", ["High School", "Diploma", "Bachelor's", "Master's", "PhD"])
institution = st.text_input("Institution Name")
field = st.text_input("Field of Study")
edu_duration = st.text_input("Duration (e.g. 2015 - 2019)")

st.subheader("💼 Work Experience")
job_title = st.text_input("Job Title")
company = st.text_input("Company Name")
work_years = st.text_input("Work Duration")
exp_years = st.number_input("Years of Experience", 0.0, 50.0, step=0.5)

st.subheader("🛠 Skills")
tech_skills = st.text_area("Technical Skills (comma separated)")
soft_skills = st.text_area("Soft Skills (comma separated)")

st.subheader("📜 Certifications & 🏆 Achievements")
certifications = st.text_area("Certifications (comma separated)")
achievements = st.text_area("Achievements or Awards (comma separated)")

st.subheader("🌐 Projects & 🗣️ Languages")
projects = st.text_area("Project Links (comma separated)")
languages = st.text_area("Languages (comma separated)")

# ---------------- Generate Button ---------------- #

if st.button("✅ Generate Resume"):
    if not name or not email or not phone:
        st.error("Please fill out name, email, and phone number.")
    else:
        doc = Document()
        user_data = {
            "name": name, "email": email, "phone": phone, "photo": photo,
            "summary": summary, "edu_level": edu_level, "institution": institution,
            "field": field, "edu_duration": edu_duration,
            "job_title": job_title, "company": company, "work_years": work_years,
            "exp_years": exp_years, "tech_skills": tech_skills, "soft_skills": soft_skills,
            "certifications": certifications, "achievements": achievements,
            "projects": projects, "languages": languages
        }

        generate_resume(doc, user_data)

        # Save DOCX
        tmp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp_docx.name)

        # Convert to PDF
        tmp_html = tmp_docx.name.replace(".docx", ".html")
        tmp_pdf = tmp_docx.name.replace(".docx", ".pdf")

        try:
            os.system(f"soffice --headless --convert-to html '{tmp_docx.name}' --outdir '{tempfile.gettempdir()}'")
            pdfkit.from_file(tmp_html, tmp_pdf)
        except:
            st.warning("⚠️ PDF generation failed. Ensure wkhtmltopdf is installed.")

        with open(tmp_docx.name, "rb") as f:
            st.download_button("⬇️ Download DOCX", f, file_name="resume.docx")

        if os.path.exists(tmp_pdf):
            with open(tmp_pdf, "rb") as pdf_file:
                st.download_button("⬇️ Download PDF", pdf_file, file_name="resume.pdf")

        st.success("🎉 Resume generated successfully!")