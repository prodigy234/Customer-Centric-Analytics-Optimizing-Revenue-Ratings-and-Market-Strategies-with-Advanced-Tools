import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import base64
import tempfile
import os
from fpdf import FPDF

st.set_page_config(page_title="AI Resume Builder", layout="centered")

st.markdown(
    """
    <style>
        .main { background-color: #111827; color: white; }
        .block-container { padding-top: 2rem; padding-bottom: 2rem; }
        .stTextInput, .stTextArea, .stSelectbox, .stMultiselect, .stFileUploader {
            background-color: #1f2937 !important;
            color: white !important;
        }
    </style>
    """,
    unsafe_allow_html=True
)

st.title("🧠 AI Resume Builder")
st.markdown("Customize your resume with professional templates and download in both **DOCX** and **PDF** formats.")

template_style = st.selectbox("Choose a Template Style", ["Classic", "Modern", "Creative"])

pdf_themes = {
    "Classic": {"font": "Times", "font_size": 12},
    "Modern": {"font": "Helvetica", "font_size": 11},
    "Creative": {"font": "Courier", "font_size": 12}
}

with st.form("resume_form", clear_on_submit=False):
    st.header("👤 Personal Information")
    name = st.text_input("Full Name")
    email = st.text_input("Email")
    phone = st.text_input("Phone Number")
    photo = st.file_uploader("Upload a Profile Photo", type=["jpg", "jpeg", "png"])

    st.header("🎯 Career Objective")
    summary = st.text_area("Professional Summary")

    st.header("📚 Education")
    edu_level = st.selectbox("Education Level", ["High School", "Diploma", "Bachelor's", "Master's", "PhD"])
    institution = st.text_input("Institution")
    field = st.text_input("Field of Study")
    edu_duration = st.text_input("Duration (e.g., 2015 - 2019)")

    st.header("💼 Experience")
    job_title = st.text_input("Job Title")
    company = st.text_input("Company Name")
    work_years = st.text_input("Years Worked")
    exp_years = st.number_input("Years of Experience", min_value=0.0, max_value=50.0, step=0.5)

    st.header("🚲 Skills")
    tech_skills = st.text_area("Technical Skills (comma separated)")
    soft_skills = st.text_area("Soft Skills (comma separated)")

    st.header("📜 Certifications & 🏆 Achievements")
    certifications = st.text_area("Certifications (comma separated)")
    achievements = st.text_area("Achievements or Awards (comma separated)")

    st.header("🌐 Projects & 🗣️ Languages")
    projects = st.text_area("Project Links (comma separated)")
    languages = st.text_area("Languages (comma separated)")

    submitted = st.form_submit_button("✅ Generate Resume")

if submitted:
    required_fields = [name.strip(), email.strip(), phone.strip()]
    if not all(required_fields):
        st.error("Please fill out **name**, **email**, and **phone number**.")
    else:
        doc = Document()

        if photo:
            image_path = os.path.join(tempfile.gettempdir(), photo.name)
            with open(image_path, "wb") as f:
                f.write(photo.read())
            doc.add_picture(image_path, width=Inches(1.5))

        heading = doc.add_heading(name, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Email: {email}")
        doc.add_paragraph(f"Phone: {phone}")
        doc.add_paragraph("")

        def add_section(title, content, bullets=True):
            if content:
                doc.add_heading(title, level=1)
                if bullets:
                    for item in content.split(","):
                        doc.add_paragraph(f"• {item.strip()}", style='List Bullet')
                else:
                    doc.add_paragraph(content)

        add_section("Career Objective", summary, bullets=False)
        if institution and field and edu_duration:
            add_section("Education", f"{edu_level} in {field} - {institution} ({edu_duration})", bullets=False)
        if job_title and company and work_years:
            add_section("Experience", f"{job_title} at {company} ({work_years}) - {exp_years} years", bullets=False)
        add_section("Technical Skills", tech_skills)
        add_section("Soft Skills", soft_skills)
        add_section("Certifications", certifications)
        add_section("Achievements", achievements)
        add_section("Projects", projects)
        add_section("Languages", languages)

        docx_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(docx_file.name)

        # -------- Generate PDF Resume --------
        pdf = FPDF()
        pdf.add_page()

        font_style = pdf_themes[template_style]["font"]
        font_size = pdf_themes[template_style]["font_size"]

        pdf.set_font(font_style, "B", 16)
        pdf.cell(200, 10, txt=name, ln=1, align="C")
        pdf.set_font(font_style, "", font_size)
        pdf.cell(200, 10, txt=f"Email: {email} | Phone: {phone}", ln=2, align="C")
        pdf.ln(10)

        def pdf_section(title, content):
            if content:
                pdf.set_font(font_style, "B", font_size)
                pdf.cell(200, 10, txt=title, ln=True)
                pdf.set_font(font_style, "", font_size)
                if isinstance(content, str):
                    lines = content.split(",") if "," in content else [content]
                    for item in lines:
                        pdf.cell(200, 8, txt=f"- {item.strip()}", ln=True)
                pdf.ln(3)

        pdf_section("Career Objective", summary)
        pdf_section("Education", f"{edu_level} in {field} - {institution} ({edu_duration})")
        pdf_section("Experience", f"{job_title} at {company} ({work_years}) - {exp_years} years")
        pdf_section("Technical Skills", tech_skills)
        pdf_section("Soft Skills", soft_skills)
        pdf_section("Certifications", certifications)
        pdf_section("Achievements", achievements)
        pdf_section("Projects", projects)
        pdf_section("Languages", languages)

        pdf_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        pdf.output(pdf_file.name)

        # -------- Download Buttons --------
        col1, col2 = st.columns(2)
        with col1:
            with open(docx_file.name, "rb") as f:
                st.download_button("📅 Download DOCX Resume", f, "resume.docx")
        with col2:
            with open(pdf_file.name, "rb") as f:
                st.download_button("📅 Download PDF Resume", f, "resume.pdf")

        st.success("✅ Resume generated successfully!")

        # Preview Section
        st.markdown(f"### 🎨 {template_style} Template Preview")
        st.write(f"**Font**: {font_style} | **Font Size**: {font_size}")
        st.info(f"This preview is styled using the {template_style} format. Open the downloaded PDF/DOCX to view the full design.")
